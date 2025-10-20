import os
import sys
import json
import requests
import io
from typing import List, Dict, Any
import openai
from pymongo import MongoClient
from pinecone import Pinecone
import numpy as np
from datetime import datetime
from dotenv import load_dotenv
from docx import Document
import re

# 加载环境变量
load_dotenv('../.env.local')

class DocumentProcessor:
    # 初始化DocumentProcessor类
    def __init__(self):
        # 获取SiliconFlow API密钥
        siliconflow_api_key = os.getenv('SILICONFLOW_API_KEY')
        if not siliconflow_api_key:
            raise ValueError("SILICONFLOW_API_KEY not found in environment variables")
        
        # 初始化SiliconFlow客户端（使用OpenAI客户端和自定义基础URL）
        self.client = openai.OpenAI(
            api_key=siliconflow_api_key,
            base_url="https://api.siliconflow.cn/v1"
        )
        
        # 初始化MongoDB
        mongo_uri = os.getenv('MONGODB_URI', 'mongodb://localhost:27017/')
        self.mongo_client = MongoClient(mongo_uri)
        self.db = self.mongo_client['allpassagent']
        self.documents_collection = self.db['documents']
        self.vectors_collection = self.db['vectors']
        
        # 初始化Pinecone (使用新版本API)
        pc = Pinecone(api_key=os.getenv('PINECONE_API_KEY'))
        self.index = pc.Index(os.getenv('PINECONE_INDEX_NAME'))
    
    # 从URL下载并解析docx文件
    def download_and_parse_docx(self, file_url: str) -> List[str]:
        """从URL下载并解析docx文件"""
        try:
            # 下载文件，禁用SSL验证并使用自定义请求头
            headers = {
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
            }
            response = requests.get(file_url, headers=headers, verify=False, timeout=30)
            response.raise_for_status()
            
            # 解析docx文件
            doc = Document(io.BytesIO(response.content))
            
            # 从所有段落中提取文本
            full_text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text += paragraph.text + "\n"
            
            # 从表格中提取文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text += cell.text + "\n"
            
            # 清理文本
            full_text = re.sub(r'\n+', '\n', full_text)
            full_text = full_text.strip()
            
            # 分割成块（每块大约500个字符）
            chunks = []
            chunk_size = 500
            overlap = 50
            
            sentences = re.split(r'[。！？\n]', full_text)
            current_chunk = ""
            
            for sentence in sentences:
                sentence = sentence.strip()
                if not sentence:
                    continue
                    
                if len(current_chunk) + len(sentence) < chunk_size:
                    current_chunk += sentence + "。"
                else:
                    if current_chunk:
                        chunks.append(current_chunk.strip())
                    current_chunk = sentence + "。"
            
            # 添加最后一个块
            if current_chunk:
                chunks.append(current_chunk.strip())
            
            return chunks
            
        except Exception as e:
            print(f"Error downloading and parsing docx: {e}")
            return []
    
    # 处理URL文档：下载、解析、生成嵌入向量并存储到Pinecone
    def process_url_document(self, file_url: str, filename: str = None) -> Dict[str, Any]:
        """ 处理URL文档：下载、解析、生成嵌入向量并存储到Pinecone"""
        try:
            # 如果未提供文件名，从URL中提取文件名
            if not filename:
                filename = file_url.split('/')[-1]
            
            # 下载并解析文档
            chunks = self.download_and_parse_docx(file_url)
            if not chunks:
                return {'error': 'Failed to download or parse document'}
            
            # 在MongoDB中创建文档记录
            document_record = {
                'filename': filename,
                'fileUrl': file_url,
                'fileType': 'docx',
                'uploadedAt': datetime.utcnow(),
                'processed': False,
                'vectorized': False,
                'chunks': chunks,
                'chunkCount': len(chunks)
            }
            
            # 将文档插入MongoDB
            result = self.documents_collection.insert_one(document_record)
            document_id = str(result.inserted_id)
            
            # 为所有块生成嵌入向量
            embeddings = self.generate_embeddings(chunks)
            if not embeddings:
                return {'error': 'Failed to generate embeddings'}
            
            # 为Pinecone准备向量数据
            vectors_to_upsert = []
            vector_records = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vector_id = f"{document_id}_chunk_{i}"
                
                # 为Pinecone准备数据
                vectors_to_upsert.append({
                    'id': vector_id,
                    'values': embedding,
                    'metadata': {
                        'doc_id': document_id,
                        'document_id': document_id,  # 保持向后兼容
                        'chunk_index': i,
                        'filename': filename,
                        'file_type': 'document',
                        'media_type': 'docx',
                        'content_type': 'text',
                        'full_content': chunk[:1000],  # 在元数据中存储前1000个字符
                        'content_summary': chunk[:200] + '...' if len(chunk) > 200 else chunk,
                        'page': i + 1,  # 添加页码信息，对于文档块使用chunk_index+1作为页码
                        'page_type': 'chunk'
                    }
                })
                
                # 为MongoDB准备数据
                vector_records.append({
                    'vector_id': vector_id,
                    'document_id': document_id,
                    'chunk_index': i,
                    'content': chunk,
                    'embedding': embedding,
                    'created_at': datetime.utcnow()
                })
            
            # 上传到Pinecone
            self.index.upsert(vectors=vectors_to_upsert)
            
            # 存储到MongoDB
            if vector_records:
                self.vectors_collection.insert_many(vector_records)
            
            # 更新文档状态
            self.documents_collection.update_one(
                {'_id': result.inserted_id},
                {
                    '$set': {
                        'vectorized': True,
                        'processed': True,
                        'processed_at': datetime.utcnow(),
                        'vector_count': len(embeddings)
                    }
                }
            )
            
            return {
                'success': True,
                'document_id': document_id,
                'filename': filename,
                'chunks_created': len(chunks),
                'vectors_created': len(embeddings),
                'message': 'Document processed successfully from URL'
            }
            
        except Exception as e:
            print(f"Error processing URL document: {e}")
            return {'error': str(e)}
        
    # 搜索相似文档：根据查询文本在Pinecone中搜索相似文档
    def search_similar_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """ 搜索相似文档：根据查询文本在Pinecone中搜索相似文档"""
        try:
            # 为查询生成嵌入向量
            query_embedding = self.generate_embeddings([query])[0]
            
            # 在Pinecone中搜索相似向量
            results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True
            )
            
            return results['matches']
            
        except Exception as e:
            print(f"Error searching similar documents: {e}")
            return []
        
    def generate_embeddings(self, texts: List[str]) -> List[List[float]]:
        """Generate embeddings for a list of texts using SiliconFlow"""
        try:
            embeddings = []
            for text in texts:
                # 截断文本以适应token限制（为安全起见，大约400个字符）
                truncated_text = text[:400] if len(text) > 400 else text
                
                response = self.client.embeddings.create(
                    model="BAAI/bge-large-zh-v1.5",
                    input=[truncated_text]
                )
                embeddings.append(response.data[0].embedding)
            
            return embeddings
        except Exception as e:
            print(f"Error generating embeddings: {e}")
            return []
    
    def process_document(self, document_id: str) -> Dict[str, Any]:
        """Process a single document: generate embeddings and store in Pinecone"""
        try:
            # 从MongoDB获取文档
            document = self.documents_collection.find_one({'_id': document_id})
            if not document:
                return {'error': f'Document {document_id} not found'}
            
            if document.get('vectorized', False):
                return {'message': f'Document {document_id} already processed'}
            
            chunks = document.get('chunks', [])
            if not chunks:
                return {'error': 'No chunks found in document'}
            
            # 为所有块生成嵌入向量
            embeddings = self.generate_embeddings(chunks)
            if not embeddings:
                return {'error': 'Failed to generate embeddings'}
            
            # 为Pinecone准备向量数据
            vectors_to_upsert = []
            vector_records = []
            
            for i, (chunk, embedding) in enumerate(zip(chunks, embeddings)):
                vector_id = f"{document_id}_chunk_{i}"
                
                # 为Pinecone准备数据
                vectors_to_upsert.append({
                    'id': vector_id,
                    'values': embedding,
                    'metadata': {
                        'doc_id': str(document_id),
                        'document_id': str(document_id),  # 保持向后兼容
                        'chunk_index': i,
                        'filename': document['filename'],
                        'file_type': 'document',
                        'media_type': document['fileType'],
                        'content_type': 'text',
                        'full_content': chunk[:1000],  # 在元数据中存储前1000个字符
                        'content_summary': chunk[:200] + '...' if len(chunk) > 200 else chunk,
                        'page': i + 1,  # 添加页码信息，对于文档块使用chunk_index+1作为页码
                        'page_type': 'chunk'
                    }
                })
                
                # 为MongoDB准备数据
                vector_records.append({
                    'vector_id': vector_id,
                    'document_id': document_id,
                    'chunk_index': i,
                    'content': chunk,
                    'embedding': embedding,
                    'created_at': datetime.utcnow()
                })
            
            # 上传到Pinecone
            self.index.upsert(vectors=vectors_to_upsert)
            
            # 存储到MongoDB
            if vector_records:
                self.vectors_collection.insert_many(vector_records)
            
            # 更新文档状态
            self.documents_collection.update_one(
                {'_id': document_id},
                {
                    '$set': {
                        'vectorized': True,
                        'processed': True,
                        'processed_at': datetime.utcnow(),
                        'vector_count': len(embeddings)
                    }
                }
            )
            
            return {
                'success': True,
                'document_id': str(document_id),
                'vectors_created': len(embeddings),
                'message': 'Document processed successfully'
            }
            
        except Exception as e:
            print(f"Error processing document {document_id}: {e}")
            return {'error': str(e)}
    
    # 搜索相似文档：根据查询文本在Pinecone中搜索相似文档
    def search_similar_documents(self, query: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Search for similar documents using vector similarity"""
        try:
            # 为查询生成嵌入向量
            query_embedding = self.generate_embeddings([query])[0]
            
            # 在Pinecone中搜索
            search_results = self.index.query(
                vector=query_embedding,
                top_k=top_k,
                include_metadata=True
            )
            
            results = []
            for match in search_results['matches']:
                results.append({
                    'id': match['id'],
                    'score': match['score'],
                    'content': match['metadata'].get('full_content', ''),
                    'filename': match['metadata'].get('filename', ''),
                    'document_id': match['metadata'].get('document_id', ''),
                    'chunk_index': match['metadata'].get('chunk_index', 0)
                })
            
            return results
            
        except Exception as e:
            print(f"Error searching documents: {e}")
            return []
    
    def get_unprocessed_documents(self) -> List[Dict[str, Any]]:
        """Get list of documents that haven't been vectorized yet"""
        try:
            cursor = self.documents_collection.find(
                {'vectorized': {'$ne': True}},
                {'_id': 1, 'filename': 1, 'uploadedAt': 1}
            )
            return list(cursor)
        except Exception as e:
            print(f"Error getting unprocessed documents: {e}")
            return []

# 主函数：从命令行处理文档
def main():
    """Main function to process documents from command line"""
    if len(sys.argv) < 2:
        print("Usage: python document_processor.py <command> [args]")
        print("Commands:")
        print("  process <document_id> - Process a specific document")
        print("  process_all - Process all unprocessed documents")
        print("  process_url <file_url> [filename] - Process document from URL")
        print("  search <query> - Search for similar content")
        return
    
    processor = DocumentProcessor()
    command = sys.argv[1]
    
    if command == 'process' and len(sys.argv) > 2:
        document_id = sys.argv[2]
        result = processor.process_document(document_id)
        print(json.dumps(result, indent=2))
    
    elif command == 'process_url' and len(sys.argv) > 2:
        file_url = sys.argv[2]
        filename = sys.argv[3] if len(sys.argv) > 3 else None
        result = processor.process_url_document(file_url, filename)
        print(json.dumps(result, indent=2))
    
    elif command == 'process_all':
        unprocessed = processor.get_unprocessed_documents()
        print(f"Found {len(unprocessed)} unprocessed documents")
        
        for doc in unprocessed:
            print(f"Processing {doc['filename']}...")
            result = processor.process_document(doc['_id'])
            print(json.dumps(result, indent=2))
    
    elif command == 'search' and len(sys.argv) > 2:
        query = ' '.join(sys.argv[2:])
        results = processor.search_similar_documents(query)
        print(json.dumps(results, indent=2))
    
    else:
        print("Invalid command or missing arguments")

if __name__ == "__main__":
    main()